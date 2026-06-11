import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re

def set_cell_background(cell, hex_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_math_run(paragraph, math_str):
    s = math_str
    replacements = [
        (r'\mathbb{R}', 'ℝ'),
        (r'\operatorname{dir}', 'dir'),
        (r'\operatorname{sign}', 'sign'),
        (r'\operatorname{clip}', 'clip'),
        (r'\operatorname{med}', 'med'),
        (r'\operatorname{MOA}', 'MOA'),
        (r'\operatorname{COG}', 'COG'),
        (r'\text{recall}', 'recall'),
        (r'\text{precision}', 'precision'),
        (r'\text{delay}', 'delay'),
        (r'\text{eff}', 'eff'),
        (r'\theta', 'θ'),
        (r'\rho', 'ρ'),
        (r'\phi', 'φ'),
        (r'\beta', 'β'),
        (r'\alpha', 'α'),
        (r'\gamma', 'γ'),
        (r'\delta', 'δ'),
        (r'\varepsilon', 'ε'),
        (r'\kappa', 'κ'),
        (r'\Omega', 'Ω'),
        (r'\mu', 'μ'),
        (r'\sigma', 'σ'),
        (r'\Phi', 'Φ'),
        (r'\in', '∈'),
        (r'\to', '→'),
        (r'\approx', '≈'),
        (r'\le', '≤'),
        (r'\ge', '≥'),
        (r'\mp', '∓'),
        (r'\pm', '±'),
        (r'\times', '×'),
        (r'\cdot', '·'),
        (r'\dots', '…'),
        (r'\lvert', '|'),
        (r'\rvert', '|'),
        (r'\,', ' '),
        (r'\{', '{'),
        (r'\}', '}'),
        (r'\ ', ' '),
        (r'\hat\beta', 'β̂'),
        (r'\hat\mu', 'μ̂'),
        (r'\hat\sigma', 'σ̂'),
        (r'\bar i', 'ī'),
        (r'\bar z', 'z̄'),
        (r'\bar u', 'ū'),
    ]
    for old, new in replacements:
        s = s.replace(old, new)
        
    i = 0
    n = len(s)
    while i < n:
        char = s[i]
        if char == '_':
            i += 1
            if i < n:
                if s[i] == '{':
                    close_idx = s.find('}', i)
                    if close_idx != -1:
                        sub_text = s[i+1:close_idx]
                        run = paragraph.add_run(sub_text)
                        run.font.subscript = True
                        run.font.italic = True
                        i = close_idx + 1
                    else:
                        run = paragraph.add_run(s[i])
                        run.font.subscript = True
                        run.font.italic = True
                        i += 1
                else:
                    run = paragraph.add_run(s[i])
                    run.font.subscript = True
                    run.font.italic = True
                    i += 1
        elif char == '^':
            i += 1
            if i < n:
                if s[i] == '{':
                    close_idx = s.find('}', i)
                    if close_idx != -1:
                        super_text = s[i+1:close_idx]
                        run = paragraph.add_run(super_text)
                        run.font.superscript = True
                        run.font.italic = True
                        i = close_idx + 1
                    else:
                        run = paragraph.add_run(s[i])
                        run.font.superscript = True
                        run.font.italic = True
                        i += 1
                else:
                    run = paragraph.add_run(s[i])
                    run.font.superscript = True
                    run.font.italic = True
                    i += 1
        else:
            run = paragraph.add_run(char)
            run.font.italic = True
            i += 1

def parse_inline_formatting(paragraph, text):
    pattern = re.compile(r'(\*\*|\*|`|\$)')
    idx = 0
    while idx < len(text):
        match = pattern.search(text, idx)
        if not match:
            paragraph.add_run(text[idx:])
            break
            
        if match.start() > idx:
            paragraph.add_run(text[idx:match.start()])
            
        token = match.group(1)
        start_pos = match.start() + len(token)
        
        end_pos = text.find(token, start_pos)
        if end_pos != -1:
            content = text[start_pos:end_pos]
            if token == '**':
                run = paragraph.add_run(content)
                run.bold = True
            elif token == '*':
                run = paragraph.add_run(content)
                run.italic = True
            elif token == '`':
                run = paragraph.add_run(content)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xA1, 0x1E, 0x51)
            elif token == '$':
                add_math_run(paragraph, content)
            idx = end_pos + len(token)
        else:
            paragraph.add_run(token)
            idx = match.start() + len(token)

if __name__ == '__main__':
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "To jest **pogrubienie**, a to *kursywa*, a to `kod` i wzor: $x_r(t) = m_r + \sigma_r z_r(t)$.")
    doc.save("scratch/test_out.docx")
    print("Done")
