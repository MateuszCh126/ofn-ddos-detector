import os
import re
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import matplotlib.pyplot as plt
from matplotlib import rcParams

# Configure matplotlib for clean math rendering
rcParams['mathtext.fontset'] = 'dejavusans'

# Styles and colors definition
COLOR_PRIMARY = "1B365D"       # Deep Navy
COLOR_SECONDARY = "4A607A"     # Muted Slate
COLOR_TEXT = "2D3748"          # Dark Charcoal
COLOR_BG_LIGHT = "F8FAFC"      # Slate White
COLOR_BORDER = "E2E8F0"        # Light Slate Gray
COLOR_CODE_CHERRY = "A11E51"   # Cherry Red for inline code

def set_cell_background(cell, hex_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_bottom_border(paragraph, color_hex="1B365D", size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{size}" w:space="6" w:color="{color_hex}"/></w:pBdr>')
    pPr.append(pBdr)

def apply_text_style(font, name="Segoe UI", size_pt=11, bold=False, italic=False, color_rgb=(0x2D, 0x37, 0x48)):
    font.name = name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(*color_rgb)

def clean_latex(s):
    """Strip unsupported sizing commands like \\big, \\Big from LaTeX."""
    s = s.replace(r'\left', r'\LEFT_TEMP')
    
    s = s.replace(r'\geq', r'\GEQ_TEMP')
    s = s.replace(r'\ge', r'\geq')
    s = s.replace(r'\GEQ_TEMP', r'\geq')
    
    s = s.replace(r'\leq', r'\LEQ_TEMP')
    s = s.replace(r'\le', r'\leq')
    s = s.replace(r'\LEQ_TEMP', r'\leq')
    
    s = s.replace(r'\LEFT_TEMP', r'\left')
    
    s = s.replace(r'\lvert', '|')
    s = s.replace(r'\rvert', '|')
    
    s = s.replace(r'\big(', '(')
    s = s.replace(r'\big)', ')')
    s = s.replace(r'\big[', '[')
    s = s.replace(r'\big]', ']')
    s = s.replace(r'\big\{', '\\{')
    s = s.replace(r'\big\}', '\\}')
    
    s = s.replace(r'\Big(', '(')
    s = s.replace(r'\Big)', ')')
    s = s.replace(r'\Big[', '[')
    s = s.replace(r'\Big]', ']')
    
    s = s.replace(r'\bigg(', '(')
    s = s.replace(r'\bigg)', ')')
    s = s.replace(r'\Bigg(', '(')
    s = s.replace(r'\Bigg)', ')')
    
    s = s.replace(r'\big', '')
    s = s.replace(r'\Big', '')
    s = s.replace(r'\bigg', '')
    s = s.replace(r'\Bigg', '')
    return s

def latex_to_png(latex_str, filepath, dpi=300):
    """Render LaTeX equation to a PNG with transparent background."""
    # Remove any newlines and replace with spaces to force math mode parsing
    latex_str = latex_str.replace('\n', ' ')
    latex_str = clean_latex(latex_str)
    
    if not latex_str.startswith('$'):
        latex_str = f"${latex_str}$"
        
    fig = plt.figure(figsize=(0.1, 0.1), facecolor='none')
    fig.set_dpi(100)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis('off')
    
    # Render text with COLOR_TEXT (2D3748) and larger fontsize for readability
    t = ax.text(0.5, 0.5, latex_str, fontsize=16, ha='center', va='center', color='#2D3748')
    
    fig.canvas.draw()
    renderer = fig.canvas.get_renderer()
    bbox = t.get_window_extent(renderer)
    
    # Bounding box is in pixels at 100 DPI
    draw_dpi = 100.0
    width_in = (bbox.width + 40) / draw_dpi
    height_in = (bbox.height + 25) / draw_dpi
    fig.set_size_inches(width_in, height_in)
    
    plt.savefig(filepath, dpi=dpi, transparent=True, bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    return width_in

def render_cases_custom(lhs, cases_val, cases_cond, filepath, dpi=300):
    """Render a LaTeX cases formula by manually laying out the components."""
    lhs = clean_latex(lhs)
    cases_val = [clean_latex(v) for v in cases_val]
    cases_cond = [clean_latex(c) for c in cases_cond]
    
    fig = plt.figure(figsize=(8, 2.2), facecolor='none')
    fig.set_dpi(100)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis('off')
    
    # 1. Render LHS
    t_lhs = ax.text(0.05, 0.5, f"${lhs}$", fontsize=16, ha='left', va='center', color='#2D3748')
    
    fig.canvas.draw()
    renderer = fig.canvas.get_renderer()
    bbox_lhs = t_lhs.get_window_extent(renderer)
    
    # Position brace next to LHS
    fig_width_px = fig.bbox.width
    lhs_right_frac = bbox_lhs.x1 / fig_width_px
    brace_x = lhs_right_frac + 0.015
    
    t_brace = ax.text(brace_x, 0.5, r"$\{$", fontsize=48, ha='left', va='center', color='#2D3748', weight='light')
    
    fig.canvas.draw()
    bbox_brace = t_brace.get_window_extent(renderer)
    brace_right_frac = bbox_brace.x1 / fig_width_px
    
    # Position case lines next to brace
    cases_x = brace_right_frac + 0.015
    
    num_cases = len(cases_val)
    if num_cases == 3:
        y_coords = [0.76, 0.50, 0.24]
    elif num_cases == 2:
        y_coords = [0.68, 0.32]
    else:
        y_coords = [0.50]
        
    t_cases = []
    for val_str, cond_str, y_coord in zip(cases_val, cases_cond, y_coords):
        case_text = f"${val_str} \\quad {cond_str}$"
        t_case = ax.text(cases_x, y_coord, case_text, fontsize=14, ha='left', va='center', color='#2D3748')
        t_cases.append(t_case)
        
    fig.canvas.draw()
    
    # Find max right coordinate
    max_x1 = 0
    for t_case in t_cases:
        bbox = t_case.get_window_extent(renderer)
        if bbox.x1 > max_x1:
            max_x1 = bbox.x1
            
    draw_dpi = 100.0
    width_in = (max_x1 + 40) / draw_dpi
    height_in = 0.5 + (num_cases * 0.45)
    
    fig.set_size_inches(width_in, height_in)
    
    plt.savefig(filepath, dpi=dpi, transparent=True, bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    return width_in

def add_math_run(paragraph, math_str):
    """Translate inline math LaTeX symbols to Unicode and add as runs."""
    s = math_str
    replacements = [
        (r'\mathbb{R}', 'ℝ'),
        (r'\operatorname{dir}', 'dir'),
        (r'\operatorname{sign}', 'sign'),
        (r'\operatorname{clip}', 'clip'),
        (r'\operatorname{med}', 'med'),
        (r'\operatorname{MOA}', 'MOA'),
        (r'\operatorname{COG}', 'COG'),
        (r'\text{recall}', 'recall'),
        (r'\text{precision}', 'precision'),
        (r'\text{delay}', 'delay'),
        (r'\text{eff}', 'eff'),
        (r'\theta', 'θ'),
        (r'\rho', 'ρ'),
        (r'\phi', 'φ'),
        (r'\beta', 'β'),
        (r'\alpha', 'α'),
        (r'\gamma', 'γ'),
        (r'\delta', 'δ'),
        (r'\varepsilon', 'ε'),
        (r'\kappa', 'κ'),
        (r'\Omega', 'Ω'),
        (r'\mu', 'μ'),
        (r'\sigma', 'σ'),
        (r'\Phi', 'Φ'),
        (r'\in', '∈'),
        (r'\to', '→'),
        (r'\approx', '≈'),
        (r'\le', '≤'),
        (r'\ge', '≥'),
        (r'\mp', '∓'),
        (r'\pm', '±'),
        (r'\times', '×'),
        (r'\cdot', '·'),
        (r'\dots', '…'),
        (r'\lvert', '|'),
        (r'\rvert', '|'),
        (r'\,', ' '),
        (r'\{', '{'),
        (r'\}', '}'),
        (r'\ ', ' '),
        (r'\hat\beta', 'β̂'),
        (r'\hat\mu', 'μ̂'),
        (r'\hat\sigma', 'σ̂'),
        (r'\bar i', 'ī'),
        (r'\bar z', 'z̄'),
        (r'\bar u', 'ū'),
    ]
    for old, new in replacements:
        s = s.replace(old, new)
        
    i = 0
    n = len(s)
    while i < n:
        char = s[i]
        if char == '_':
            i += 1
            if i < n:
                if s[i] == '{':
                    close_idx = s.find('}', i)
                    if close_idx != -1:
                        sub_text = s[i+1:close_idx]
                        run = paragraph.add_run(sub_text)
                        apply_text_style(run.font, size_pt=9.5, italic=True)
                        run.font.subscript = True
                        i = close_idx + 1
                    else:
                        run = paragraph.add_run(s[i])
                        apply_text_style(run.font, size_pt=9.5, italic=True)
                        run.font.subscript = True
                        i += 1
                else:
                    run = paragraph.add_run(s[i])
                    apply_text_style(run.font, size_pt=9.5, italic=True)
                    run.font.subscript = True
                    i += 1
        elif char == '^':
            i += 1
            if i < n:
                if s[i] == '{':
                    close_idx = s.find('}', i)
                    if close_idx != -1:
                        super_text = s[i+1:close_idx]
                        run = paragraph.add_run(super_text)
                        apply_text_style(run.font, size_pt=9.5, italic=True)
                        run.font.superscript = True
                        i = close_idx + 1
                    else:
                        run = paragraph.add_run(s[i])
                        apply_text_style(run.font, size_pt=9.5, italic=True)
                        run.font.superscript = True
                        i += 1
                else:
                    run = paragraph.add_run(s[i])
                    apply_text_style(run.font, size_pt=9.5, italic=True)
                    run.font.superscript = True
                    i += 1
        else:
            run = paragraph.add_run(char)
            apply_text_style(run.font, size_pt=11, italic=True)
            i += 1

def add_formatted_text(paragraph, text, size_pt=11, is_list=False):
    """Parse bold, italic, code, and inline math formatting and add runs to paragraph."""
    pattern = re.compile(r'(\*\*|\*|`|\$)')
    idx = 0
    while idx < len(text):
        match = pattern.search(text, idx)
        if not match:
            # Plain text run
            run = paragraph.add_run(text[idx:])
            apply_text_style(run.font, size_pt=size_pt)
            break
            
        if match.start() > idx:
            run = paragraph.add_run(text[idx:match.start()])
            apply_text_style(run.font, size_pt=size_pt)
            
        token = match.group(1)
        start_pos = match.start() + len(token)
        
        end_pos = text.find(token, start_pos)
        if end_pos != -1:
            content = text[start_pos:end_pos]
            if token == '**':
                run = paragraph.add_run(content)
                apply_text_style(run.font, size_pt=size_pt, bold=True)
            elif token == '*':
                run = paragraph.add_run(content)
                apply_text_style(run.font, size_pt=size_pt, italic=True)
            elif token == '`':
                run = paragraph.add_run(content)
                apply_text_style(run.font, name='Consolas', size_pt=size_pt - 1.5, color_rgb=(161, 30, 81))
            elif token == '$':
                add_math_run(paragraph, content)
            idx = end_pos + len(token)
        else:
            # Unmatched delimiter, treat as plain text
            run = paragraph.add_run(token)
            apply_text_style(run.font, size_pt=size_pt)
            idx = match.start() + len(token)

def parse_markdown(md_text):
    """Parse Markdown file into blocks."""
    lines = md_text.split('\n')
    blocks = []
    i = 0
    n = len(lines)
    
    while i < n:
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Header 1
        if line.startswith('# '):
            blocks.append(('h1', line[2:].strip()))
            i += 1
            continue
            
        # Header 2
        if line.startswith('## '):
            blocks.append(('h2', line[3:].strip()))
            i += 1
            continue
            
        # Header 3
        if line.startswith('### '):
            blocks.append(('h3', line[4:].strip()))
            i += 1
            continue
            
        # Horizontal rule
        if line == '---':
            blocks.append(('hr', ''))
            i += 1
            continue
            
        # Block math
        if line.startswith('$$'):
            math_lines = []
            if line.endswith('$$') and len(line) > 2:
                math_lines.append(line[2:-2].strip())
            else:
                math_lines.append(line[2:].strip())
                i += 1
                while i < n and not lines[i].strip().endswith('$$'):
                    math_lines.append(lines[i].strip())
                    i += 1
                if i < n:
                    end_line = lines[i].strip()
                    if end_line != '$$':
                        math_lines.append(end_line[:-2].strip())
            blocks.append(('block_math', '\n'.join(math_lines)))
            i += 1
            continue
            
        # Table
        if line.startswith('|'):
            table_rows = []
            while i < n and lines[i].strip().startswith('|'):
                table_row_line = lines[i].strip()
                if not all(c in '|- :+' for c in table_row_line):
                    cols = [col.strip() for col in table_row_line.split('|')[1:-1]]
                    table_rows.append(cols)
                i += 1
            blocks.append(('table', table_rows))
            continue
            
        # Bullet list item
        if line.startswith('- ') or line.startswith('* '):
            blocks.append(('bullet_list_item', line[2:].strip()))
            i += 1
            continue
            
        # Numbered list item
        num_list_match = re.match(r'^(\d+)\.\s+(.*)$', line)
        if num_list_match:
            num = num_list_match.group(1)
            item_text = num_list_match.group(2)
            blocks.append(('numbered_list_item', (num, item_text)))
            i += 1
            continue
            
        # Standard paragraph
        paragraph_lines = [line]
        i += 1
        while i < n:
            next_line = lines[i].strip()
            if not next_line:
                break
            if (next_line.startswith('#') or 
                next_line.startswith('$$') or 
                next_line.startswith('|') or 
                next_line.startswith('- ') or 
                next_line.startswith('* ') or 
                next_line == '---' or 
                re.match(r'^(\d+)\.\s+(.*)$', next_line)):
                break
            paragraph_lines.append(next_line)
            i += 1
        blocks.append(('paragraph', ' '.join(paragraph_lines)))
        
    return blocks

def main():
    # Read the markdown documentation
    md_path = "D:/github/github/ofn-ddos-detector/DOKUMENTACJA_NAUKOWA.md"
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
        
    blocks = parse_markdown(md_content)
    
    # Initialize Document
    doc = Document()
    
    # Configure Margins: 1 inch on all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Setup temp directory for math PNGs
    os.makedirs('scratch/temp_math', exist_ok=True)
    
    # Set default paragraph style properties
    style_normal = doc.styles['Normal']
    apply_text_style(style_normal.font, name='Segoe UI', size_pt=11, color_rgb=(0x2D, 0x37, 0x48))
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(8)
    style_normal.paragraph_format.space_before = Pt(0)
    
    # Track states
    is_abstract_mode = False
    
    # Convert blocks
    block_math_index = 0
    
    for b_idx, (b_type, b_content) in enumerate(blocks):
        # Determine if we are writing the abstract (Streszczenie)
        if b_type == 'h2' and b_content == 'Streszczenie':
            is_abstract_mode = True
            # Create a 1x1 table for callout box for the abstract
            abstract_table = doc.add_table(rows=1, cols=1)
            abstract_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            abstract_table.autofit = False
            abstract_table.columns[0].width = Inches(6.5)
            
            cell = abstract_table.cell(0, 0)
            cell.width = Inches(6.5)
            set_cell_background(cell, COLOR_BG_LIGHT)
            set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
            
            # Left border in deep navy, other borders none
            tcPr = cell._tc.get_or_add_tcPr()
            borders_xml = (
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="none"/>'
                f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{COLOR_PRIMARY}"/>'
                f'<w:bottom w:val="none"/>'
                f'<w:right w:val="none"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(parse_xml(borders_xml))
            
            # Title inside abstract box
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("Streszczenie")
            apply_text_style(run.font, name='Segoe UI Semibold', size_pt=12.5, bold=True, color_rgb=(0x1B, 0x36, 0x5D))
            continue
            
        if is_abstract_mode:
            if b_type == 'paragraph':
                # Add abstract content inside the callout box paragraph
                cell = abstract_table.cell(0, 0)
                p = cell.add_paragraph()
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                # Format text with size_pt=10.5 and italic inside callout
                add_formatted_text(p, b_content, size_pt=10.5)
                # Apply italic styling to abstract runs
                for r in p.runs:
                    r.italic = True
                continue
            elif b_type == 'hr':
                # The horizontal line ends the abstract mode
                is_abstract_mode = False
                # Add a normal empty spacer after the abstract callout box
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(12)
                continue
                
        # Normal (Non-abstract) elements
        if b_idx == 0 and b_type == 'h1':
            # Document Title Page Block
            # Title paragraph
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(0)
            p_title.paragraph_format.space_after = Pt(12)
            p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p_title.add_run(b_content)
            apply_text_style(run.font, name='Segoe UI', size_pt=24, bold=True, color_rgb=(0x1B, 0x36, 0x5D))
            
            # Check if second block is subtitle
            if len(blocks) > 1 and blocks[1][0] == 'paragraph':
                # Subtitle paragraph
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.space_before = Pt(0)
                p_sub.paragraph_format.space_after = Pt(18)
                p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Extract text inside **Dokumentacja...** and parse
                subtitle_text = blocks[1][1]
                add_formatted_text(p_sub, subtitle_text, size_pt=12)
                for r in p_sub.runs:
                    r.font.italic = True
                    r.font.color.rgb = RGBColor(0x4A, 0x60, 0x7A)
                
                # Bottom border line on subtitle
                add_bottom_border(p_sub, color_hex=COLOR_PRIMARY, size="16") # 2 pt bottom accent line
                
            continue
            
        if b_idx == 1 and b_type == 'paragraph':
            # Handled by title section, skip
            continue
            
        if b_type == 'h1':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(22)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(b_content)
            apply_text_style(run.font, name='Segoe UI Semibold', size_pt=16, bold=True, color_rgb=(0x1B, 0x36, 0x5D))
            add_bottom_border(p, color_hex=COLOR_PRIMARY, size="8") # 1 pt border
            
        elif b_type == 'h2':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(b_content)
            apply_text_style(run.font, name='Segoe UI Semibold', size_pt=13, bold=True, color_rgb=(0x1B, 0x36, 0x5D))
            
        elif b_type == 'h3':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(b_content)
            apply_text_style(run.font, name='Segoe UI Semibold', size_pt=11.5, bold=True, color_rgb=(0x4A, 0x60, 0x7A))
            
        elif b_type == 'paragraph':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(8)
            add_formatted_text(p, b_content)
            
        elif b_type == 'bullet_list_item':
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.25)
            add_formatted_text(p, b_content)
            
        elif b_type == 'numbered_list_item':
            num, item_text = b_content
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.25)
            add_formatted_text(p, item_text)
            
        elif b_type == 'block_math':
            block_math_index += 1
            formula = b_content
            
            # Check if this is a boxed formula (Formula 5)
            # Markdown: \boxed{\,A - A = (0, 0) = \mathbf{0}\,}
            is_boxed = False
            boxed_match = re.match(r'\\boxed\{(.*)\}', formula.strip())
            if boxed_match:
                is_boxed = True
                inner_formula = boxed_match.group(1).strip()
                if inner_formula.startswith(r'\,'):
                    inner_formula = inner_formula[2:]
                if inner_formula.endswith(r'\,'):
                    inner_formula = inner_formula[:-2]
                formula = inner_formula
                
            filepath = f"scratch/temp_math/math_{block_math_index}.png"
            
            # Check if this is a cases formula
            if r'\begin{cases}' in formula:
                start_idx = formula.find(r'\begin{cases}')
                end_idx = formula.find(r'\end{cases}')
                lhs = formula[:start_idx].strip()
                cases_content = formula[start_idx + len(r'\begin{cases}'):end_idx].strip()
                trailing_math = formula[end_idx + len(r'\end{cases}'):].strip()
                
                raw_lines = re.split(r'\\\\(?:\[.*?\])?', cases_content)
                cases_val = []
                cases_cond = []
                for line in raw_lines:
                    line = line.strip()
                    if not line:
                        continue
                    parts = line.split('&')
                    if len(parts) == 2:
                        cases_val.append(parts[0].strip())
                        cases_cond.append(parts[1].strip())
                    else:
                        cases_val.append(line)
                        cases_cond.append("")
                        
                if trailing_math:
                    if cases_cond:
                        if not trailing_math.startswith(r'\qquad') and not trailing_math.startswith(r'\quad'):
                            cases_cond[-1] = cases_cond[-1] + r'\quad ' + trailing_math
                        else:
                            cases_cond[-1] = cases_cond[-1] + ' ' + trailing_math
                            
                width_in = render_cases_custom(lhs, cases_val, cases_cond, filepath)
            else:
                width_in = latex_to_png(formula, filepath)
            
            # Insert into Word
            if is_boxed:
                # Place inside a beautiful 1x1 table (callout box)
                box_table = doc.add_table(rows=1, cols=1)
                box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                box_table.autofit = False
                box_table.columns[0].width = Inches(6.5)
                
                cell = box_table.cell(0, 0)
                cell.width = Inches(6.5)
                set_cell_background(cell, COLOR_BG_LIGHT)
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                
                # Single-line navy border around the boxed equation callout box
                tcPr = cell._tc.get_or_add_tcPr()
                borders_xml = (
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{COLOR_PRIMARY}"/>'
                    f'<w:left w:val="single" w:sz="12" w:space="0" w:color="{COLOR_PRIMARY}"/>'
                    f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{COLOR_PRIMARY}"/>'
                    f'<w:right w:val="single" w:sz="6" w:space="0" w:color="{COLOR_PRIMARY}"/>'
                    f'</w:tcBorders>'
                )
                tcPr.append(parse_xml(borders_xml))
                
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                
                run = p.add_run()
                run.add_picture(filepath, width=Inches(min(width_in, 6.0)))
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(8)
                
                run = p.add_run()
                run.add_picture(filepath, width=Inches(min(width_in, 6.2)))
                
        elif b_type == 'table':
            rows = b_content
            if len(rows) == 0:
                continue
                
            num_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            
            # Apply sleek borders to the table XML
            tblPr = table._tbl.tblPr
            borders_xml = (
                f'<w:tblBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
                f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="{COLOR_PRIMARY}"/>'
                f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
                f'<w:left w:val="none"/>'
                f'<w:right w:val="none"/>'
                f'<w:insideV w:val="none"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(parse_xml(borders_xml))
            
            # Fill rows
            for r_idx, row_data in enumerate(rows):
                for c_idx, val in enumerate(row_data):
                    cell = table.cell(r_idx, c_idx)
                    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    
                    if r_idx == 0:
                        # Header Row Styling
                        set_cell_background(cell, COLOR_PRIMARY)
                        # Center header text
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Remove markdown bold wrappers inside text if any
                        clean_val = val.replace('**', '')
                        run = p.add_run(clean_val)
                        apply_text_style(run.font, name='Segoe UI Semibold', size_pt=10.5, bold=True, color_rgb=(255, 255, 255))
                    else:
                        # Data Row Styling
                        if r_idx % 2 == 0:
                            # Zebra stripe background
                            set_cell_background(cell, COLOR_BG_LIGHT)
                        
                        # Set alignment: if column 0, left aligned; else center or left depending on contents
                        if c_idx == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                        add_formatted_text(p, val, size_pt=10)
                        
            # Add a normal paragraph spacer after the table
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(8)
            
        elif b_type == 'hr':
            # Non-abstract HR is represented as a divider paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_bottom_border(p, color_hex="CBD5E1", size="4") # thin divider line
            
    # Save the output
    out_docx_path = "D:/github/github/ofn-ddos-detector/DOKUMENTACJA_NAUKOWA.docx"
    doc.save(out_docx_path)
    print(f"Successfully generated beautiful docx: {out_docx_path}")

if __name__ == '__main__':
    main()
